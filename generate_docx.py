import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_logbook_docx():
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    # 1. Header
    p_uni = doc.add_paragraph()
    p_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_uni = p_uni.add_run("ADVENTIST UNIVERSITY OF CENTRAL AFRICA\n")
    run_uni.bold = True
    run_uni.font.size = Pt(16)
    run_uni.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    run_sub = p_uni.add_run("P.O. Box 2461 Kigali, Rwanda | www.auca.ac.rw | info@auca.ac.rw\n")
    run_sub.font.size = Pt(9.5)
    run_sub.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    run_fac = p_uni.add_run("FACULTY OF INFORMATION TECHNOLOGY\n")
    run_fac.bold = True
    run_fac.font.size = Pt(13)
    run_fac.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_doc_title = p_title.add_run("STUDENTS' ATTACHMENT LOGBOOK & PRACTICUM REPORT")
    run_doc_title.bold = True
    run_doc_title.font.size = Pt(14)
    run_doc_title.underline = True

    # 2. Form I Details
    h1 = doc.add_heading("FORM I: PERSONAL & PLACEMENT DETAILS", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    h1.runs[0].font.size = Pt(12)

    table_info = doc.add_table(rows=6, cols=2)
    table_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    data_info = [
        ("Student Name:", "Regine PACIS", "Faculty:", "Faculty of Information Technology"),
        ("Student ID:", "______________________", "Department:", "Software Engineering / Information Technology"),
        ("Course Code:", "INAT 353 Industrial Attachment", "Academic Year:", "2025 / 2026"),
        ("Host Organization:", "RoboScope Engineering Project", "Host Department:", "Full-Stack & Cloud Systems Engineering"),
        ("Project Title:", "RoboScope Robotics & Embodied AI Platform", "Location / Address:", "Kigali, Rwanda"),
        ("Industry Supervisor:", "Lead Software Engineer", "Supervisor Title:", "Senior Systems Architect"),
    ]
    
    for row_idx, data_row in enumerate(data_info):
        row = table_info.rows[row_idx]
        cell_a = row.cells[0]
        cell_b = row.cells[1]
        cell_a.text = f"{data_row[0]} {data_row[1]}"
        cell_b.text = f"{data_row[2]} {data_row[3]}"
        set_cell_shading(cell_a, "F8FAFC")
        set_cell_shading(cell_b, "F8FAFC")
        set_cell_margins(cell_a, 60, 60, 100, 100)
        set_cell_margins(cell_b, 60, 60, 100, 100)

    doc.add_paragraph()

    # 3. Form III Week 1
    doc.add_page_break()
    h2 = doc.add_heading("LOG FORM III: DAILY SUMMARY REPORT (WEEK 1)", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    h2.runs[0].font.size = Pt(12)

    p_w1 = doc.add_paragraph()
    p_w1.add_run("Week 1 Focus: Codebase Audit, Diagnostics, Dataset Pipelines & Video Sync\n").bold = True

    t_w1 = doc.add_table(rows=1, cols=6)
    t_w1.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = t_w1.rows[0].cells
    headers = ["Day", "Brief Description of Work / Activity Performed", "Time In/Out", "Hours", "Lessons Learnt", "Challenges Faced"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_shading(hdr_cells[i], "E2E8F0")
        set_cell_margins(hdr_cells[i], 80, 80, 80, 80)

    week1_rows = [
        ("Mon", "Codebase Audit & Package Restructuring: Audited full-stack architecture; corrected Python package initializations (_init_.py to __init__.py); added missing dependencies (pandas, pydantic); validated SQLite schema and FTS5 triggers.", "08:00\n17:00", "8 hrs", "PEP-8 standards and package namespace isolation in modular FastAPI architectures.", "Module import failures due to malformed init files."),
        ("Tue", "Robotics Dataset Ingestion & Async I/O: Refactored Hugging Face LeRobot downloader; resolved shadowed API route handlers in datasets.py; implemented FastAPI BackgroundTasks; built dynamic Parquet trajectory discovery.", "08:00\n17:00", "8 hrs", "Asynchronous job queuing prevents HTTP worker thread starvation during multi-GB downloads.", "Non-standard Parquet episode numbering causing file-not-found errors."),
        ("Wed", "Video Streaming & API Consolidation: Configured multi-camera MP4 video streams; eliminated duplicate/broken frontend API clients into a unified api.js; fixed missing React useState hook in useVideoSync.js.", "08:00\n17:00", "8 hrs", "Centralized API clients and sub-millisecond video synchronization using requestVideoFrameCallback.", "API query param mismatches (camera vs key) and runtime React hook crashes."),
        ("Thu", "Canvas Trajectory Visualizer with uPlot: Developed TrajectoryPlot.jsx with ResizeObserver; implemented custom canvas vertical cursor scrubber synchronized with real-time video playback and frame stepping.", "08:00\n17:00", "8 hrs", "High-performance canvas rendering for multi-channel robotic sensor time-series telemetry.", "Timeline scrubber coordinate mapping under dynamic viewport resizing."),
        ("Fri", "Catalog & Live Feed Integration: Integrated automated arXiv robotics crawler (cs.RO, cs.LG); built interactive TagPicker.jsx with autocomplete; implemented live download progress polling in DatasetCatalog.jsx.", "08:00\n17:00", "8 hrs", "Reactive polling patterns for real-time monitoring of asynchronous background cloud jobs.", "State race conditions during concurrent dataset downloads and UI re-renders.")
    ]

    for row_data in week1_rows:
        row = t_w1.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_margins(cell, 60, 60, 60, 60)

    # Week 1 Total
    row_tot1 = t_w1.add_row()
    row_tot1.cells[0].text = "TOTAL"
    row_tot1.cells[1].text = "Week 1 Cumulative Total: 40 Working Hours Completed"
    row_tot1.cells[3].text = "40 hrs"
    for cell in row_tot1.cells:
        set_cell_shading(cell, "F1F5F9")
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.bold = True

    # 4. Form III Week 2
    doc.add_page_break()
    h3 = doc.add_heading("LOG FORM III: DAILY SUMMARY REPORT (WEEK 2)", level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    h3.runs[0].font.size = Pt(12)

    p_w2 = doc.add_paragraph()
    p_w2.add_run("Week 2 Focus: Security Engineering, Auth UI, Dashboard & Production Cloud Deployment\n").bold = True

    t_w2 = doc.add_table(rows=1, cols=6)
    t_w2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2_cells = t_w2.rows[0].cells
    for i, h in enumerate(headers):
        hdr2_cells[i].text = h
        hdr2_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr2_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_shading(hdr2_cells[i], "E2E8F0")
        set_cell_margins(hdr2_cells[i], 80, 80, 80, 80)

    week2_rows = [
        ("Mon", "Security Architecture & Cryptographic Engine: Migrated SQLite database schema to support users and user-linked saved_items; developed backend/auth.py using PBKDF2-HMAC-SHA256 (100k rounds) and signed JWT tokens.", "08:00\n17:00", "8 hrs", "Industry-standard cryptographic hashing and stateless JWT authorization without external C-compilers.", "Ensuring zero-compilation cross-platform portability across Windows, WSL, and Linux containers."),
        ("Tue", "Auth API & Frontend State Management: Built routers/auth.py (/register, /login, /me); created React AuthContext.jsx for persistent JWT sessions; created glassmorphic AuthPage.jsx with validation.", "08:00\n17:00", "8 hrs", "Token persistence via localStorage, HTTP Authorization header injection, and client-side route guards.", "Handling token expiration gracefully and maintaining synchronized state across page reloads."),
        ("Wed", "Personalized Research Dashboard Development: Engineered GET /api/dashboard aggregation endpoint; built Dashboard.jsx with live KPI cards, bookmark manager, dataset launcher, and category progress bars.", "08:00\n17:00", "8 hrs", "Architecting analytical dashboard interfaces with high information density and intuitive visual hierarchy.", "Efficiently computing relational SQLite metric aggregations across multiple tables in a single query."),
        ("Thu", "Design System Overhaul & Build Verification: Upgraded App.css to a sleek cybernetic dark aesthetic; added user avatar profile badge and dropdown menu; fixed optional Rollup binaries for Windows/WSL; verified production build.", "08:00\n17:00", "8 hrs", "CSS design tokens, responsive grid layouts, and asset bundle optimization for Single Page Applications.", "Cross-platform native binary conflicts between Windows host and WSL environment."),
        ("Fri", "Containerization & Production Cloud Deployment: Configured static SPA serving in FastAPI; created multi-stage Dockerfile, build.sh, and render.yaml; initialized Git, resolved nested submodules, and deployed live to Render.", "08:00\n17:00", "8 hrs", "CI/CD deployment pipelines, SPA client-side routing fallback on Nginx/Uvicorn, and zero-cost cloud PaaS setup.", "Nested .git repository blocking staging; configuring GitHub Personal Access Token authentication.")
    ]

    for row_data in week2_rows:
        row = t_w2.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_margins(cell, 60, 60, 60, 60)

    # Week 2 Total
    row_tot2 = t_w2.add_row()
    row_tot2.cells[0].text = "TOTAL"
    row_tot2.cells[1].text = "Cumulative Attachment Total: 80 Working Hours Completed"
    row_tot2.cells[3].text = "40 hrs"
    for cell in row_tot2.cells:
        set_cell_shading(cell, "F1F5F9")
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.bold = True

    # 5. Form II Sample
    doc.add_page_break()
    h4 = doc.add_heading("LOG FORM II: DAILY DETAILED DESCRIPTION OF WORK (SAMPLE ENTRY)", level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    h4.runs[0].font.size = Pt(12)

    t_f2 = doc.add_table(rows=7, cols=2)
    t_f2.alignment = WD_TABLE_ALIGNMENT.CENTER
    form2_data = [
        ("Sequence of Operation for Job", "1. Formulated REST endpoint specifications for user dashboard aggregation in FastAPI.\n2. Built relational SQL queries in SQLite computing user bookmarked papers, ready datasets, taxonomy tags, and category breakdown.\n3. Constructed React Dashboard.jsx page with responsive 4-column KPI cards and interactive list views.\n4. Implemented real-time progress bars visualizing paper distribution across cs.RO (Robotics) and cs.LG (Machine Learning).\n5. Integrated route protection redirecting unauthenticated visitors to /login."),
        ("Tools & Equipment Used", "Python 3.11, FastAPI, SQLite3, React 18, React Router v7, Visual Studio Code / Antigravity IDE, Git."),
        ("Tasks Completed", "• Completed backend/routers/dashboard.py with GET /api/dashboard.\n• Completed frontend/src/pages/Dashboard.jsx with KPI cards and dataset quick-launcher.\n• Connected optimistic UI bookmark removal and manual arXiv sync trigger."),
        ("Tasks in Progress", "End-to-end production build testing and cross-platform native binary packaging."),
        ("Next Day's Tasks", "Overhaul global CSS design system and execute multi-platform build verification."),
        ("Problems / Challenges", "Structuring responsive CSS Grid layouts that dynamically adapt multi-series progress bars across mobile and desktop viewports."),
        ("Student's Recommendations", "Use dedicated analytics aggregation endpoints rather than multiple round-trip client requests to optimize frontend render latency.")
    ]

    for idx, (label, content) in enumerate(form2_data):
        row = t_f2.rows[idx]
        cell_lbl = row.cells[0]
        cell_cnt = row.cells[1]
        cell_lbl.text = label
        cell_cnt.text = content
        cell_lbl.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell_lbl, "F1F5F9")
        set_cell_margins(cell_lbl, 80, 80, 80, 80)
        set_cell_margins(cell_cnt, 80, 80, 80, 80)

    # 6. SWOT Analysis
    doc.add_page_break()
    h5 = doc.add_heading("SECTION XX: SELF EVALUATION REPORT (SWOT ANALYSIS)", level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    h5.runs[0].font.size = Pt(12)

    doc.add_heading("1. STRENGTHS", level=2).runs[0].font.color.rgb = RGBColor(0x16, 0x65, 0x34)
    doc.add_paragraph("• Full-Stack Integration: Successfully integrated reactive React 18 frontend with asynchronous Python FastAPI backend via REST APIs and JWT security.\n• Problem Solving: Isolated and resolved complex bugs (shadowed route handlers, package misconfigurations, and video-telemetry sync drift).\n• DevOps Mastery: Gained practical expertise in Docker containerization, Nginx reverse proxying, and automated cloud PaaS deployment on Render.")

    doc.add_heading("2. WEAKNESSES", level=2).runs[0].font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
    doc.add_paragraph("• Initial Cross-Platform Tooling Familiarity: Encountered platform-specific native binary conflicts between Windows and WSL Linux environments during early build attempts.")

    doc.add_heading("3. OPPORTUNITIES", level=2).runs[0].font.color.rgb = RGBColor(0x07, 0x59, 0x85)
    doc.add_paragraph("• Embodied AI & Robotics: Hands-on exposure to Hugging Face LeRobot datasets and PyArrow Parquet formats opened career avenues in AI data engineering.\n• Cloud Architecture: Practical experience in zero-cost cloud deployments provides high readiness for software engineering roles.")

    doc.add_heading("4. THREATS & CHALLENGES OVERCOME", level=2).runs[0].font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    doc.add_paragraph("• Git Submodules & Authentication: Overcame nested repository blockers and GitHub Personal Access Token authentication requirements to deploy the production codebase.")

    doc.add_heading("5. CONCLUSION & LESSONS LEARNT", level=2).runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    doc.add_paragraph("The industrial attachment provided intensive, real-world full-stack development and DevOps experience. Beyond writing clean code, the exercise reinforced the necessity of structured software architecture, database indexing, user security engineering, and cloud deployment automation. The completed RoboScope platform stands as a production-grade, highly responsive system ready for academic and industrial research use.")

    # Signatures
    p_sign = doc.add_paragraph("\n\n")
    table_sign = doc.add_table(rows=1, cols=3)
    table_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    s_cells = table_sign.rows[0].cells
    s_cells[0].text = "Student Signature:\n\n_______________________\nDate: _________________"
    s_cells[1].text = "Field Supervisor Signature:\n\n_______________________\nDate: _________________"
    s_cells[2].text = "University Supervisor Signature:\n\n_______________________\nDate: _________________"

    output_path = "c:/Users/LENOVO/Desktop/robscope/AUCA_STUDENTS_ATTACHMENT_LOGBOOK.docx"
    doc.save(output_path)
    print(f"Successfully generated DOCX at: {output_path}")

if __name__ == "__main__":
    create_logbook_docx()
